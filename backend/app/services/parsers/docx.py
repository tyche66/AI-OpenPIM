from __future__ import annotations

import asyncio
import hashlib
from io import BytesIO

from docx import Document

from app.services.parsers.base import ParserError, ParseResult, ParserProtocol


class DocxTextParser(ParserProtocol):
    """Extract paragraphs and table text from DOCX files."""

    name = "python-docx"
    version = "1.1.2"

    async def parse(self, file_bytes: bytes, *, file_name: str = "") -> ParseResult:
        return await asyncio.to_thread(self._parse_sync, file_bytes)

    def _parse_sync(self, file_bytes: bytes) -> ParseResult:
        try:
            document = Document(BytesIO(file_bytes))
        except Exception as exc:  # noqa: BLE001 - keep parser errors controlled
            raise ParserError("parse_failed", "DOCX 文本解析失败") from exc

        parts: list[str] = []
        parts.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts).strip()
        if not text:
            raise ParserError("empty_document", "DOCX 未包含可提取文本")
        return ParseResult(
            text=text,
            content_hash=hashlib.sha256(text.encode("utf-8")).hexdigest(),
            parser_name=self.name,
            parser_version=self.version,
            page_count=None,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
        )
